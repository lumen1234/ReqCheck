import os
import sys
from docx import Document

# 测试直接读取docx文件的原始内容
def test_docx_raw():
    print("测试直接读取docx文件的原始内容...")
    filepath = "MEMS陀螺软件需求规格说明.docx"
    
    if not os.path.exists(filepath):
        print(f"文件不存在: {filepath}")
        return
    
    # 直接读取docx文件
    document = Document(filepath)
    
    print(f"文档共 {len(document.paragraphs)} 段")
    print("\n前100段的原始内容:")
    
    for i, para in enumerate(document.paragraphs[:100]):
        text = para.text
        style_name = para.style.name
        print(f"[{i+1}] 样式: {style_name}, 内容: '{text}'")

if __name__ == "__main__":
    test_docx_raw()
